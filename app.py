import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict
import re

# -----------------------------------------------------------------------------
# إعدادات الواجهة الرسومية والتنسيق الصارم (CSS)
# -----------------------------------------------------------------------------
st.set_page_config(page_title="نظام فرز المناطق الصارم والمطابق", layout="wide", page_icon="⚖️")
st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; }
    div.stButton > button { background-color: #2C3E50; color: white; width: 100%; font-weight: bold; border-radius: 8px; padding: 10px;}
    div.stButton > button:hover { background-color: #1A252F; color: #F1C40F;}
    .report-box { background-color: #ECF0F1; padding: 12px; border-radius: 8px; border-right: 5px solid #2C3E50; text-align: right; margin-bottom: 10px;}
    .vessel-box { background-color: #FDEDEC; padding: 12px; border-radius: 8px; border-right: 5px solid #CB4335; text-align: right; margin-bottom: 10px;}
    .stat-title { font-size: 13px; color: #7F8C8D; font-weight: bold; }
    .stat-value { font-size: 18px; color: #2C3E50; font-weight: bold; margin-top: 3px;}
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: right;'>نظام الفرز الصارم - مع فلتر التدميج العائلي المتسلسل 🎯⚖️</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: right;'>تم ضبط النظام لدمج الأفراد المنفصلين داخل كتل عوائلهم بالمنطقة المستهدفة تلقائياً بحسب الأبوة، الجد، واللقب.</p>", unsafe_allow_html=True)

LIST_OF_ZONES = [
    "المنطقة الأولى", "المنطقة الثانية", "المنطقة الثالثة", 
    "المنطقة الرابعة", "المنطقة الخامسة", "المنطقة السادسة", "المنطقة السابعة"
]
VESSEL_ZONE = "الوعاء الأصلي (المتبقي)"

# -----------------------------------------------------------------------------
# دوال التنسيق والهيكلة المستندية (Word)
# -----------------------------------------------------------------------------
def set_cell_direction(cell, direction='btLr'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)
    tcPr.append(textDirection)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567))) 
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def apply_rtl(doc):
    for style in doc.styles:
        if hasattr(style, 'font'): style.font.rtl = True

def preprocess_name(name):
    name = re.sub(r'\s+', ' ', name.strip())
    for prefix in ['عبد ', 'ابو ', 'أبو ', 'أم ', 'ام ', 'امة ', 'أمة ', 'آل ']:
        name = name.replace(prefix, prefix.strip() + '_')
    return name

def is_female(name):
    first_word = name.strip().split()[0]
    if first_word in ['ام', 'أم', 'امة', 'أمة', 'آمال']: return True
    male_exceptions = ['حمزة', 'اسامة', 'أسامة', 'حذيفة', 'قتادة', 'طلحة', 'خليفة', 'معاوية', 'عطية', 'حارثة', 'عروة', 'عبيدة', 'ميسرة', 'سلامة']
    if first_word.endswith('ة') and first_word not in male_exceptions: return True
    female_names = [
        'زينب', 'مريم', 'شهد', 'نور', 'هدى', 'ندى', 'ليلى', 'سها', 'مها', 'ريم', 'سعاد', 'هند',
        'ايمان', 'إيمان', 'رحاب', 'سحر', 'سمر', 'كوثر', 'غدير', 'حنان', 'منى', 'اسماء', 'أسماء',
        'اسراء', 'إسراء', 'شيماء', 'دعاء', 'وفاء', 'رجاء', 'سناء', 'لقاء', 'نجلاء', 'حوراء',
        'عذراء', 'زهراء', 'بيداء', 'رفل', 'رسل', 'ضحى', 'شمس', 'قمر', 'فرح', 'مرح', 'ابتسام',
        'إبتسام', 'امال', 'ابتهال', 'انتصار', 'إنتصار', 'انعام', 'أنعام', 'بشرى', 'ذكرى', 'رؤى', 
        'سرى', 'سروة', 'سهام', 'شروق', 'صبا', 'عواطف', 'فاتن', 'لمياء', 'لينا', 'محاسن', 'مروة', 
        'نوال', 'هاجر', 'وجدان', 'ياسمين', 'يسرى', 'بنين', 'غفران', 'فردوس', 'افراح', 'أفراح', 
        'اخلاص', 'إخلاص', 'ازهار', 'أزهار', 'الحان', 'ألحان', 'انوار', 'أنوار', 'اشواق', 'أشواق', 
        'براء', 'تبارك', 'تغريد', 'جمانة', 'جنان', 'جواهر', 'حلا', 'خلود', 'داليا', 'دلال', 'رغد', 
        'رنا', 'رنين', 'روان', 'سالي', 'سجى', 'سوزان', 'صابرين', 'عالية', 'عبير', 'عفاف', 'غادة', 
        'ليال', 'مي', 'ميسون', 'نادية', 'نجاة', 'نغم', 'نهى', 'هديل', 'هناء', 'هيفاء', 'ورود', 
        'وسن', 'وعد', 'يقين', 'إلهام', 'الهام', 'تهاني', 'سلوى', 'رشا', 'سهير', 'منال', 'آية', 
        'اية', 'غسق', 'شيرين', 'نسرين', 'جيهان', 'إيناس', 'ايناس', 'رواء'
    ]
    if first_word in female_names: return True
    return False

# -----------------------------------------------------------------------------
# دالة استخراج وتفكيك الاسم بالتفصيل للمطابقة المتسلسلة الصارمة
# -----------------------------------------------------------------------------
def get_name_fallback_parts(name_str):
    clean_name = re.sub(r'\s+', ' ', name_str.strip()).replace('_', ' ')
    words = clean_name.split()
    father = ""
    grand = ""
    title = ""
    if len(words) >= 2:
        title = words[-1]
    if len(words) >= 3:
        father = words[1]
    if len(words) >= 4:
        grand = words[2]
    return {"father": father, "grand": grand, "title": title}

# -----------------------------------------------------------------------------
# محرك الربط الرصين وتجميع الأصول والفروع 
# -----------------------------------------------------------------------------
def extract_and_group_data(file_obj):
    doc_in = Document(file_obj)
    records = []
    
    for table in doc_in.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not any(cells) or "المركز" in "".join(cells) or "الوكيل" in "".join(cells) or "اسم رب" in "".join(cells): continue
            
            name_idx = -1; max_len = 0
            for i, c in enumerate(cells):
                if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
                    if len(c) > max_len: max_len = len(c); name_idx = i
            if name_idx == -1: continue
            
            card_indices = [i for i, c in enumerate(cells) if c.isdigit() and len(c) >= 5]
            if not card_indices: continue
            card_num = cells[card_indices[1]] if len(card_indices) >= 2 else cells[card_indices[0]]
            old_card = cells[card_indices[0]] if len(card_indices) >= 2 else ""
            seq = next((cells[i] for i in range(len(cells)-1, card_indices[-1], -1) if cells[i].isdigit()), "-")
            digit_cells = [cells[i] for i in range(name_idx) if cells[i].isdigit()]
            
            if len(digit_cells) >= 3: withheld, eligible, total = digit_cells[0], digit_cells[1], digit_cells[2]
            elif len(digit_cells) == 2: withheld, eligible, total = "0", digit_cells[0], digit_cells[1]
            else: withheld, eligible, total = "0", "0", "0"
                
            records.append({
                'withheld': str(withheld), 'eligible': str(eligible), 'total': str(total),
                'name': cells[name_idx].strip(), 'old_card': old_card, 'card': card_num, 'seq': seq
            })

    if not records: return None, None

    male_records = []
    female_records = []
    
    for rec in records:
        rec['processed_name'] = preprocess_name(rec['name'])
        rec['is_female'] = is_female(rec['name'])
        if rec['is_female']:
            female_records.append(rec)
        else:
            words = rec['processed_name'].split()
            rec['father_string'] = " ".join(words[1:]) if len(words) > 1 else rec['processed_name']
            male_records.append(rec)

    male_full_names = {r['processed_name']: r for r in male_records}
    unique_fathers = list(set(r['father_string'] for r in male_records))
    
    family_map = {}
    for f in unique_fathers:
        if f in male_full_names:
            family_map[f] = f
        else:
            matches = [other for other in unique_fathers if other.startswith(f + " ") and other != f]
            family_map[f] = matches[0] if len(matches) == 1 else f

    final_family_map = {}
    for f in unique_fathers:
        current = f
        for _ in range(5):
            target = family_map.get(current, current)
            if target == current: break
            current = target
        final_family_map[f] = current

    final_family_groups = defaultdict(list)
    for rec in male_records:
        if rec['processed_name'] in unique_fathers:
            assigned_root = final_family_map[rec['processed_name']]
        else:
            assigned_root = final_family_map[rec['father_string']]
        rec['root_id'] = assigned_root
        final_family_groups[assigned_root].append(rec)

    families_list = []
    for root, members in final_family_groups.items():
        def sort_key(m):
            is_father = 0 if m['processed_name'] == root else 1
            return (is_father, m['name'])
            
        members_sorted = sorted(members, key=sort_key)
        families_list.append({
            'root_id': root,
            'father_display': members_sorted[0]['processed_name'].replace('_', ' '),
            'leading_name': members_sorted[0]['name'],
            'size': len(members_sorted),
            'members': members_sorted,
            'is_female_group': False
        })

    for rec in female_records:
        fam_id = 'female_' + str(id(rec))
        rec['root_id'] = fam_id
        families_list.append({
            'root_id': fam_id,
            'father_display': rec['name'],
            'leading_name': rec['name'],
            'size': 1,
            'members': [rec],
            'is_female_group': True
        })
        
    return families_list, records

def create_word_file(sorted_records):
    doc_out = Document()
    apply_rtl(doc_out)
    headers = ['ت', 'رقم البطاقة', 'رقم البطاقة القديم', 'اسم رب الاسرة', 'الافراد الكلية', 'الافراد المستحقة', 'الافراد المحجوبين']
    table = doc_out.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[i].paragraphs[0].runs: run.font.size, run.font.bold, run.font.name = Pt(13), True, 'Arial'
            
    set_cell_direction(hdr_cells[4], 'btLr'); set_cell_direction(hdr_cells[5], 'btLr'); set_cell_direction(hdr_cells[6], 'btLr')
    
    for idx, rec in enumerate(sorted_records):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = rec['card']
        row_cells[2].text = rec['old_card']
        row_cells[3].text = rec['name']
        row_cells[4].text = rec['total']
        row_cells[5].text = rec['eligible']
        row_cells[6].text = rec['withheld']
        for i in range(7):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in row_cells[i].paragraphs[0].runs: run.font.size, run.font.name = Pt(15), 'Arial'
                
    widths = [1.2, 3.5, 3.5, 5.0, 0.8, 0.8, 0.8]
    for row in table.rows:
        for idx, width in enumerate(widths): set_cell_width(row.cells[idx], width)
            
    buffer = BytesIO()
    doc_out.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# إدارة الحالة والواجهة التفاعلية (الوعاء التناقضي والتصفية والفلاتر)
# -----------------------------------------------------------------------------
st.markdown("<h3 style='text-align: right;'>📂 خطوة 1: رفع ملف البيانات الأصلي</h3>", unsafe_allow_html=True)
uploaded_file = st.file_uploader("", type=['docx'], label_visibility="collapsed")

if uploaded_file:
    if 'vessel_file' not in st.session_state or st.session_state.vessel_file != uploaded_file.name:
        families, all_records = extract_and_group_data(uploaded_file)
        if families:
            st.session_state.families = families
            st.session_state.all_records = all_records
            st.session_state.vessel_file = uploaded_file.name
            st.session_state.family_zones = {fam['root_id']: VESSEL_ZONE for fam in families}

    if 'families' in st.session_state:
        available_names = []
        vessel_count = 0
        
        for fam in st.session_state.families:
            if st.session_state.family_zones[fam['root_id']] == VESSEL_ZONE:
                vessel_count += len(fam['members'])
                for member in fam['members']:
                    available_names.append(member['name'])
        
        available_names = sorted(list(set(available_names)))

        st.markdown("<br><hr>", unsafe_allow_html=True)
        st.markdown(f"<h3 style='text-align: right;'>⚡ خطوة 2: لوحة الفرز بنظام الوعاء المتناقص (المتبقي حالياً: {vessel_count} فرد)</h3>", unsafe_allow_html=True)
        
        if len(available_names) > 0:
            col_name, col_zone, col_action = st.columns([4, 3, 3])
            
            with col_name:
                selected_person = st.selectbox("🔎 ابحث عن اسم (العوائل الموزعة تختفي تلقائياً من هنا):", options=available_names, key="search_box")
                
            with col_zone:
                chosen_zone = st.selectbox("🏢 اختر المنطقة المستهدفة (من 1 إلى 7):", options=LIST_OF_ZONES)
                
            with col_action:
                st.markdown("<div style='padding-top: 28px;'></div>", unsafe_allow_html=True)
                if st.button("🚀 سحب العائلة وإخراجها من الوعاء"):
                    matched_record = next((r for r in st.session_state.all_records if r['name'] == selected_person), None)
                    if matched_record:
                        target_root = matched_record['root_id']
                        st.session_state.family_zones[target_root] = chosen_zone
                        bros_count = len([r for r in st.session_state.all_records if r['root_id'] == target_root])
                        st.success(f"✅ تم سحب العائلة بنجاح! نُقِل {bros_count} أفراد إلى {chosen_zone} واختفوا من الوعاء.")
                        st.rerun()
        else:
            st.success("🎉 مبروك! تم إفراغ الوعاء الأصلي بالكامل وتوزيع كافة العوائل على المناطق السبعة.")

        # ---------------------------------------------------------------------
        # توليد وحفظ الملفات الثمانية النهائية مع تطبيق فلتر الدمج الهرمي
        # ---------------------------------------------------------------------
        st.markdown("<br><hr>", unsafe_allow_html=True)
        if st.button("⚙️ إنتاج وتثبيت مستندات الـ 7 مناطق + ملف المتبقي"):
            
            zone_fams = {zone: [] for zone in LIST_OF_ZONES + [VESSEL_ZONE]}
            for fam in st.session_state.families:
                current_assigned = st.session_state.family_zones[fam['root_id']]
                zone_fams[current_assigned].append(fam)
            
            def finalize_zone_with_fallback(families_in_zone, is_vessel=False):
                # إذا كنا في وعاء المتبقي، لا ندمج الفروع لأنهم لم يوزعوا أصلاً
                if is_vessel:
                    males = [f for f in families_in_zone if not f['is_female_group']]
                    females = [f for f in families_in_zone if f['is_female_group']]
                    sorted_males = sorted(males, key=lambda x: x['leading_name'])
                    sorted_females = sorted(females, key=lambda x: x['leading_name'])
                    final_list = []
                    for f in sorted_males: final_list.extend(f['members'])
                    for f in sorted_females: final_list.extend(f['members'])
                    return final_list

                # فصل الكتل العائلية الكبيرة (الأساسية) عن الأفراد المنفصلين (المرشحين للدمج)
                primary_families = []
                standalone_candidates = []
                
                for f in families_in_zone:
                    if f['is_female_group'] or f['size'] == 1:
                        standalone_candidates.append(f)
                    else:
                        primary_families.append(f)
                
                # إذا كانت المنطقة لا تحتوي إلا على أفراد منفصلين، نجعل الرجال منهم كتل أساسية مؤقتاً لدمج النساء معهم
                if not primary_families:
                    primary_families = [f for f in families_in_zone if not f['is_female_group']]
                    standalone_candidates = [f for f in families_in_zone if f['is_female_group']]

                # استخراج أجزاء الاسم لكتل العوائل الأساسية للاستعداد للمطابقة
                for fam in primary_families:
                    fam['parts'] = get_name_fallback_parts(fam['leading_name'])
                    fam['injected_members'] = list(fam['members']) # نسخة للتعديل

                unmatched_standalone = []

                # تطبيق فلتر التدميج المتسلسل الصارم للأفراد المنفصلين
                for cand in standalone_candidates:
                    cand_member = cand['members'][0]
                    cand_parts = get_name_fallback_parts(cand_member['name'])
                    matched_fam = None

                    # الفلتر 1: تطابق (اسم الأب + الجد + اللقب)
                    if cand_parts['father'] and cand_parts['grand'] and cand_parts['title']:
                        for fam in primary_families:
                            fp = fam['parts']
                            if fp['father'] == cand_parts['father'] and fp['grand'] == cand_parts['grand'] and fp['title'] == cand_parts['title']:
                                matched_fam = fam
                                break

                    # الفلتر 2: تطابق (الجد + اللقب)
                    if not matched_fam and cand_parts['grand'] and cand_parts['title']:
                        for fam in primary_families:
                            fp = fam['parts']
                            if fp['grand'] == cand_parts['grand'] and fp['title'] == cand_parts['title']:
                                matched_fam = fam
                                break

                    # الفلتر 3: تطابق (اللقب)
                    if not matched_fam and cand_parts['title']:
                        for fam in primary_families:
                            fp = fam['parts']
                            if fp['title'] == cand_parts['title']:
                                matched_fam = fam
                                break

                    # إذا عثر الفلتر المتسلسل على عائلته، يتم حشره معهم فوراً في الجدول
                    if matched_fam:
                        matched_fam['injected_members'].append(cand_member)
                    else:
                        # إذا لم يطابق أي تسلسل، يبقى منفرداً
                        unmatched_standalone.append(cand)

                # بناء القائمة النهائية للمنطقة وترتيبها رصيناً
                sorted_primaries = sorted(primary_families, key=lambda x: (x['leading_name'][0] if x['leading_name'] else 'أ', -len(x['injected_members']), x['leading_name']))
                
                final_list = []
                for fam in sorted_primaries:
                    final_list.extend(fam['injected_members'])

                # عزل الأشخاص الذين لم يطابقوا أي عائلة (الرجال أولاً، ثم النساء في نهاية القائمة تماماً)
                unmatched_males = [f for f in unmatched_standalone if not f['is_female_group']]
                unmatched_females = [f for f in unmatched_standalone if f['is_female_group']]
                
                for f in sorted(unmatched_males, key=lambda x: x['leading_name']): final_list.extend(f['members'])
                for f in sorted(unmatched_females, key=lambda x: x['leading_name']): final_list.extend(f['members'])
                
                return final_list

            final_buffers = {}
            final_counts = {}
            
            # معالجة وحفظ الملفات الثمانية
            for zone in LIST_OF_ZONES + [VESSEL_ZONE]:
                is_vessel_flag = (zone == VESSEL_ZONE)
                zone_records = finalize_zone_with_fallback(zone_fams[zone], is_vessel=is_vessel_flag)
                final_counts[zone] = len(zone_records)
                final_buffers[zone] = create_word_file(zone_records) if zone_records else None

            st.markdown("<h3 style='text-align: right;'>📥 خطوة 3: تحميل الملفات المفروزة والمطابقة نهائياً</h3>", unsafe_allow_html=True)
            
            st.markdown("<h4 style='text-align: right; color:#CB4335;'>📦 الوعاء الأصلي (الملف الثامن للمتبقي):</h4>", unsafe_allow_html=True)
            col_vessel = st.columns(1)[0]
            with col_vessel:
                st.markdown(f"<div class='vessel-box'><div class='stat-title'>ملف المتبقي الشامل</div><div class='stat-value'>{final_counts[VESSEL_ZONE]} فرد لم يتم سحبهم بعد</div></div>", unsafe_allow_html=True)
                if final_buffers[VESSEL_ZONE]:
                    st.download_button("📥 تحميل ملف المتبقي الأصلي.docx", data=final_buffers[VESSEL_ZONE], file_name="ملف_المتبقي_الأصلي.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else: st.info("الوعاء فارغ بالكامل، كل العوائل وزعت بنجاح!")

            st.markdown("<br><h4 style='text-align: right; color:#2C3E50;'>🏢 مستندات المناطق السبعة المستقلة (المدمجة هرمياً):</h4>", unsafe_allow_html=True)
            
            r1_cols = st.columns(4)
            for i in range(4):
                zone_name = LIST_OF_ZONES[i]
                with r1_cols[i]:
                    st.markdown(f"<div class='report-box'><div class='stat-title'>{zone_name}</div><div class='stat-value'>{final_counts[zone_name]} فرد قيد جاهز</div></div>", unsafe_allow_html=True)
                    if final_buffers[zone_name]:
                        st.download_button(f"📥 تحميل {zone_name}.docx", data=final_buffers[zone_name], file_name=f"{zone_name.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else: st.warning("المنطقة فارغة")

            st.markdown("<br>", unsafe_allow_html=True)
            
            r2_cols = st.columns(3)
            for i in range(3):
                zone_name = LIST_OF_ZONES[4 + i]
                with r2_cols[i]:
                    st.markdown(f"<div class='report-box'><div class='stat-title'>{zone_name}</div><div class='stat-value'>{final_counts[zone_name]} فرد قيد جاهز</div></div>", unsafe_allow_html=True)
                    if final_buffers[zone_name]:
                        st.download_button(f"📥 تحميل {zone_name}.docx", data=final_buffers[zone_name], file_name=f"{zone_name.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else: st.warning("المنطقة فارغة")
                    
            st.success("🎉 اكتملت المعالجة الصارمة! الفلتر حصر جميع الأسماء المنفصلة مع عوائلهم داخل الجداول بامتياز.")
else:
    st.info("👋 يرجى رفع ملف الوورد الأساسي لتفعيل لوحة الفرز بنظام الوعاء المتناقص والتدميج الهرمي.")
