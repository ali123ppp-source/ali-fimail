# main.py
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import word_formatter

def preprocess_name(name):
    """معالجة الأسماء المركبة مؤقتاً لتجنب أخطاء الفرز مثل (عبد الله، ابو الفضل)"""
    name = name.replace('عبد ', 'عبد_').replace('ابو ', 'ابو_').replace('ام ', 'ام_').replace('امة ', 'امة_')
    return name

def postprocess_name(name):
    """إعادة الأسماء المركبة لوضعها الطبيعي بعد الفرز"""
    return name.replace('_', ' ')

def process_and_generate(input_path, output_path):
    print("[*] جاري قراءة البيانات...")
    # 1. استخراج البيانات من ملف الوورد
    doc_in = Document(input_path)
    lines = [para.text.strip() for para in doc_in.paragraphs if para.text.strip()]
    
    records = []
    for line in lines:
        parts = line.split(',')
        if len(parts) == 7 and parts[0].isdigit() and parts[-1].isdigit():
            records.append({
                'withheld': parts[0].strip(),
                'eligible': parts[1].strip(),
                'total': parts[2].strip(),
                'name': parts[3].strip(),
                'old_card': parts[4].strip(),
                'card': parts[5].strip(),
                'seq': parts[6].strip()
            })
            
    print("[*] جاري بناء شجرة العوائل وفرزها...")
    # 2. بناء شجرة العائلة وفرزها
    names_set = {r['name'] for r in records}
    children_map = {name: [] for name in names_set}
    parents_map = {}
    
    for name in names_set:
        processed_name = preprocess_name(name)
        words = processed_name.split()
        for i in range(1, len(words) - 1):
            potential_parent_processed = " ".join(words[i:])
            potential_parent = postprocess_name(potential_parent_processed)
            if potential_parent in names_set:
                children_map[potential_parent].append(name)
                parents_map[name] = potential_parent
                break
                
    # إيجاد الآباء وترتيبهم أبجدياً
    roots = [name for name in names_set if name not in parents_map]
    roots.sort()
    
    record_by_name = {r['name']: r for r in records}
    sorted_records = []
    
    def add_family(person):
        sorted_records.append(record_by_name[person])
        children_map[person].sort() # ترتيب الأبناء أبجدياً
        for child in children_map[person]:
            add_family(child)
            
    for root in roots:
        add_family(root)
        
    print("[*] جاري توليد وتنسيق ملف الوورد النهائي...")
    # 3. إنشاء ملف الوورد وتنسيقه
    doc_out = Document()
    word_formatter.apply_rtl(doc_out)
    
    headers = ['ت', 'رقم البطاقة', 'رقم البطاقة القديم', 'اسم رب الاسرة', 'الافراد الكلية', 'الافراد المستحقة', 'الافراد المحجوبين']
    table = doc_out.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.autofit = False
    
    # تنسيق العناوين (حجم 14 وخط غامق)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Arial'
            
    # تدوير الحقول الثلاثة المتغيرة 90 درجة بشكل عمودي
    word_formatter.set_cell_direction(hdr_cells[4], 'btLr')
    word_formatter.set_cell_direction(hdr_cells[5], 'btLr')
    word_formatter.set_cell_direction(hdr_cells[6], 'btLr')
    
    # تعبئة البيانات (حجم 16)
    for idx, rec in enumerate(sorted_records):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = rec['card']
        row_cells[2].text = rec['old_card']
        row_cells[3].text = rec['name']
        row_cells[4].text = rec['total']
        row_cells[5].text = rec['eligible']
        row_cells[6].text = rec['withheld']
        
        for i in range(7):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(16)
                run.font.name = 'Arial'
                
    # تحديد مقاسات الأعمدة بالسنتيمتر (حقل رب الأسرة = 5 سم، والأعمدة العمودية أضيق ما يمكن)
    widths = [1.2, 3.5, 3.5, 5.0, 0.8, 0.8, 0.8]
    for row in table.rows:
        for idx, width in enumerate(widths):
            word_formatter.set_cell_width(row.cells[idx], width)
            
    doc_out.save(output_path)
    print(f"[+] تم تصدير الملف بنجاح إلى: {output_path}")

if __name__ == '__main__':
    # المسار الخاص بالملف المراد ترتيبه، ومسار الإخراج
    INPUT_FILE = "data.docx"
    OUTPUT_FILE = "sorted_output.docx"
    
    # لتشغيل النظام
    process_and_generate(INPUT_FILE, OUTPUT_FILE)